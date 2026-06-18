"""
OSMnx Routing Service
A FastAPI microservice for accurate street network routing using OSMnx.

Based on: https://geoffboeing.com/2016/11/osmnx-python-street-networks/

Features:
- Downloads and caches street networks from OpenStreetMap
- Calculates shortest paths for driving, walking, biking
- Returns accurate distances and travel times
- Supports batch route calculations
- PDF text extraction using PyMuPDF
"""

import os
import io
import logging
import base64
from typing import Optional, List, Literal
from contextlib import asynccontextmanager

import osmnx as ox
import networkx as nx
import fitz  # PyMuPDF
import docx  # python-docx
from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from cachetools import TTLCache
from diskcache import Cache

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configure OSMnx
ox.settings.use_cache = True
ox.settings.cache_folder = "./osmnx_cache"
ox.settings.log_console = True

# In-memory cache for graphs (TTL: 1 hour, max 50 graphs)
graph_cache = TTLCache(maxsize=50, ttl=3600)

# Disk cache for persistent storage
disk_cache = Cache("./route_cache")


# Pydantic Models
class Coordinates(BaseModel):
    lat: float = Field(..., ge=-90, le=90)
    lng: float = Field(..., ge=-180, le=180)


class RouteRequest(BaseModel):
    origin: Coordinates
    destination: Coordinates
    mode: Literal["drive", "walk", "bike"] = "drive"


class BatchRouteRequest(BaseModel):
    routes: List[RouteRequest]


class RouteResponse(BaseModel):
    distance_km: float
    duration_minutes: float
    mode: str
    path_coordinates: Optional[List[List[float]]] = None
    success: bool = True
    error: Optional[str] = None


class BatchRouteResponse(BaseModel):
    routes: List[RouteResponse]
    total_distance_km: float
    total_duration_minutes: float


class HealthResponse(BaseModel):
    status: str
    cache_size: int
    osmnx_version: str


# Helper functions
def get_network_type(mode: str) -> str:
    """Map mode to OSMnx network type."""
    mapping = {
        "drive": "drive",
        "walk": "walk", 
        "bike": "bike"
    }
    return mapping.get(mode, "drive")


def get_speed_kmh(mode: str) -> float:
    """Get average speed for travel time estimation."""
    speeds = {
        "drive": 40,  # Average urban driving speed
        "walk": 5,    # Walking speed
        "bike": 15    # Cycling speed
    }
    return speeds.get(mode, 40)


def get_graph_key(lat: float, lng: float, mode: str, dist: int = 5000) -> str:
    """Generate cache key for graph."""
    # Round coordinates to create larger cache regions
    lat_rounded = round(lat, 2)
    lng_rounded = round(lng, 2)
    return f"{lat_rounded}_{lng_rounded}_{mode}_{dist}"


async def get_or_download_graph(lat: float, lng: float, mode: str, dist: int = 5000):
    """Get graph from cache or download from OSM."""
    cache_key = get_graph_key(lat, lng, mode, dist)
    
    # Check memory cache
    if cache_key in graph_cache:
        logger.info(f"Graph cache hit: {cache_key}")
        return graph_cache[cache_key]
    
    # Check disk cache
    if cache_key in disk_cache:
        logger.info(f"Disk cache hit: {cache_key}")
        G = disk_cache[cache_key]
        graph_cache[cache_key] = G
        return G
    
    # Download from OSM
    logger.info(f"Downloading graph for {lat}, {lng}, mode={mode}, dist={dist}")
    try:
        network_type = get_network_type(mode)
        G = ox.graph_from_point(
            (lat, lng),
            dist=dist,
            network_type=network_type,
            simplify=True
        )
        
        # Add speeds and travel times
        G = ox.speed.add_edge_speeds(G)
        G = ox.speed.add_edge_travel_times(G)
        
        # Cache the graph
        graph_cache[cache_key] = G
        disk_cache[cache_key] = G
        
        logger.info(f"Downloaded and cached graph: {len(G.nodes)} nodes, {len(G.edges)} edges")
        return G
    except Exception as e:
        logger.error(f"Failed to download graph: {e}")
        raise


def calculate_route(G, origin: Coordinates, destination: Coordinates, mode: str) -> RouteResponse:
    """Calculate shortest path between two points."""
    try:
        # Find nearest nodes to origin and destination
        orig_node = ox.nearest_nodes(G, origin.lng, origin.lat)
        dest_node = ox.nearest_nodes(G, destination.lng, destination.lat)
        
        # Calculate shortest path by travel time
        try:
            route = ox.shortest_path(G, orig_node, dest_node, weight="travel_time")
        except nx.NetworkXNoPath:
            # Fallback to length-based routing
            route = ox.shortest_path(G, orig_node, dest_node, weight="length")
        
        if route is None:
            return RouteResponse(
                distance_km=0,
                duration_minutes=0,
                mode=mode,
                success=False,
                error="No path found between points"
            )
        
        # Calculate total distance and travel time
        edge_lengths = ox.routing.route_to_gdf(G, route)["length"].sum()
        distance_km = edge_lengths / 1000
        
        # Calculate travel time
        try:
            edge_times = ox.routing.route_to_gdf(G, route)["travel_time"].sum()
            duration_minutes = edge_times / 60
        except KeyError:
            # Fallback: estimate based on distance and average speed
            duration_minutes = (distance_km / get_speed_kmh(mode)) * 60
        
        # Get route coordinates for visualization
        route_coords = [[G.nodes[node]["x"], G.nodes[node]["y"]] for node in route]
        
        return RouteResponse(
            distance_km=round(distance_km, 2),
            duration_minutes=round(duration_minutes, 1),
            mode=mode,
            path_coordinates=route_coords,
            success=True
        )
        
    except Exception as e:
        logger.error(f"Route calculation error: {e}")
        return RouteResponse(
            distance_km=0,
            duration_minutes=0,
            mode=mode,
            success=False,
            error=str(e)
        )


# Haversine fallback for when OSMnx fails
def haversine_distance(origin: Coordinates, destination: Coordinates) -> float:
    """Calculate great-circle distance as fallback."""
    from math import radians, sin, cos, sqrt, atan2
    
    R = 6371  # Earth's radius in km
    
    lat1, lng1 = radians(origin.lat), radians(origin.lng)
    lat2, lng2 = radians(destination.lat), radians(destination.lng)
    
    dlat = lat2 - lat1
    dlng = lng2 - lng1
    
    a = sin(dlat/2)**2 + cos(lat1) * cos(lat2) * sin(dlng/2)**2
    c = 2 * atan2(sqrt(a), sqrt(1-a))
    
    return R * c


# FastAPI App
@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup and shutdown events."""
    logger.info("Starting OSMnx Routing Service")
    logger.info(f"OSMnx version: {ox.__version__}")
    
    # Create cache directories
    os.makedirs("./osmnx_cache", exist_ok=True)
    os.makedirs("./route_cache", exist_ok=True)
    
    yield
    
    logger.info("Shutting down OSMnx Routing Service")
    disk_cache.close()


app = FastAPI(
    title="OSMnx Routing Service",
    description="Accurate street network routing using OSMnx and OpenStreetMap",
    version="1.0.0",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint."""
    return HealthResponse(
        status="healthy",
        cache_size=len(graph_cache),
        osmnx_version=ox.__version__
    )


@app.post("/route", response_model=RouteResponse)
async def calculate_single_route(request: RouteRequest):
    """
    Calculate route between two points.
    
    Uses OSMnx to download street network and calculate shortest path.
    Falls back to Haversine distance if routing fails.
    """
    try:
        # Calculate midpoint for graph download
        mid_lat = (request.origin.lat + request.destination.lat) / 2
        mid_lng = (request.origin.lng + request.destination.lng) / 2
        
        # Calculate required graph radius (with buffer)
        direct_distance = haversine_distance(request.origin, request.destination)
        graph_radius = max(int(direct_distance * 1500), 5000)  # At least 5km
        graph_radius = min(graph_radius, 50000)  # Cap at 50km
        
        # Get or download the graph
        G = await get_or_download_graph(mid_lat, mid_lng, request.mode, graph_radius)
        
        # Calculate route
        result = calculate_route(G, request.origin, request.destination, request.mode)
        
        # If OSMnx routing failed, use fallback
        if not result.success:
            direct_km = haversine_distance(request.origin, request.destination)
            # Apply road factor (roads are ~1.4x direct distance on average)
            road_km = direct_km * 1.4
            speed = get_speed_kmh(request.mode)
            
            return RouteResponse(
                distance_km=round(road_km, 2),
                duration_minutes=round((road_km / speed) * 60, 1),
                mode=request.mode,
                success=True,
                error="Used fallback calculation"
            )
        
        return result
        
    except Exception as e:
        logger.error(f"Route calculation failed: {e}")
        
        # Fallback to Haversine
        direct_km = haversine_distance(request.origin, request.destination)
        road_km = direct_km * 1.4
        speed = get_speed_kmh(request.mode)
        
        return RouteResponse(
            distance_km=round(road_km, 2),
            duration_minutes=round((road_km / speed) * 60, 1),
            mode=request.mode,
            success=True,
            error=f"Used fallback: {str(e)}"
        )


@app.post("/routes/batch", response_model=BatchRouteResponse)
async def calculate_batch_routes(request: BatchRouteRequest):
    """
    Calculate multiple routes in batch.
    
    More efficient than individual calls as it can reuse cached graphs.
    """
    results = []
    total_distance = 0
    total_duration = 0
    
    for route_request in request.routes:
        result = await calculate_single_route(route_request)
        results.append(result)
        
        if result.success:
            total_distance += result.distance_km
            total_duration += result.duration_minutes
    
    return BatchRouteResponse(
        routes=results,
        total_distance_km=round(total_distance, 2),
        total_duration_minutes=round(total_duration, 1)
    )


@app.post("/preload")
async def preload_graph(lat: float, lng: float, mode: str = "drive", radius: int = 10000):
    """
    Pre-load a graph for a region.
    
    Useful for warming up the cache before routing requests.
    """
    try:
        G = await get_or_download_graph(lat, lng, mode, radius)
        return {
            "status": "success",
            "nodes": len(G.nodes),
            "edges": len(G.edges),
            "mode": mode,
            "radius": radius
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============ PDF Extraction ============

class DocumentExtractRequest(BaseModel):
    """Request model for base64 document extraction."""
    data_base64: str
    filename: Optional[str] = "document"
    file_type: Literal["pdf", "docx"] = "pdf"


class DocumentExtractResponse(BaseModel):
    """Response model for document extraction."""
    text: str
    page_count: int
    success: bool = True
    error: Optional[str] = None


# Keep old models for backward compatibility
class PDFExtractRequest(BaseModel):
    """Request model for base64 PDF extraction."""
    pdf_base64: str
    filename: Optional[str] = "document.pdf"


class PDFExtractResponse(BaseModel):
    """Response model for PDF extraction."""
    text: str
    page_count: int
    success: bool = True
    error: Optional[str] = None


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> tuple[str, int]:
    """
    Extract text from PDF bytes using PyMuPDF.
    Fast extraction optimized for multi-language documents (EN, CN, TH).
    
    Returns:
        tuple: (extracted_text, page_count)
    """
    text_parts = []
    
    # Open PDF from bytes
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page_count = len(doc)
    
    for page_num, page in enumerate(doc):
        # Extract text with better unicode support for Chinese/Thai
        page_text = page.get_text("text", flags=fitz.TEXT_PRESERVE_WHITESPACE)
        if page_text.strip():
            text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
    
    doc.close()
    
    return "\n\n".join(text_parts), page_count


def extract_text_from_docx_bytes(docx_bytes: bytes) -> tuple[str, int]:
    """
    Extract text from DOCX bytes using python-docx.
    Fast extraction optimized for multi-language documents (EN, CN, TH).
    
    Returns:
        tuple: (extracted_text, paragraph_count)
    """
    text_parts = []
    
    # Open DOCX from bytes
    doc = docx.Document(io.BytesIO(docx_bytes))
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts), len(doc.paragraphs)


@app.post("/extract-document", response_model=DocumentExtractResponse)
async def extract_document_text(request: DocumentExtractRequest):
    """
    Extract text from a base64-encoded document (PDF or DOCX).
    
    Fast text extraction optimized for English, Chinese, and Thai.
    """
    try:
        logger.info(f"Extracting text from {request.file_type}: {request.filename}")
        
        # Decode base64
        try:
            doc_bytes = base64.b64decode(request.data_base64)
        except Exception as e:
            return DocumentExtractResponse(
                text="",
                page_count=0,
                success=False,
                error=f"Invalid base64 data: {str(e)}"
            )
        
        # Extract text based on file type
        if request.file_type == "pdf":
            text, page_count = extract_text_from_pdf_bytes(doc_bytes)
        elif request.file_type == "docx":
            text, page_count = extract_text_from_docx_bytes(doc_bytes)
        else:
            return DocumentExtractResponse(
                text="",
                page_count=0,
                success=False,
                error=f"Unsupported file type: {request.file_type}"
            )
        
        logger.info(f"Extracted {len(text)} characters from {page_count} pages/paragraphs")
        
        return DocumentExtractResponse(
            text=text,
            page_count=page_count,
            success=True
        )
        
    except Exception as e:
        logger.error(f"Document extraction error: {e}")
        return DocumentExtractResponse(
            text="",
            page_count=0,
            success=False,
            error=str(e)
        )


@app.post("/extract-pdf", response_model=PDFExtractResponse)
async def extract_pdf_text(request: PDFExtractRequest):
    """
    Extract text from a base64-encoded PDF.
    
    Fast text extraction using PyMuPDF (fitz).
    """
    try:
        logger.info(f"Extracting text from PDF: {request.filename}")
        
        # Decode base64
        try:
            pdf_bytes = base64.b64decode(request.pdf_base64)
        except Exception as e:
            return PDFExtractResponse(
                text="",
                page_count=0,
                success=False,
                error=f"Invalid base64 data: {str(e)}"
            )
        
        # Extract text
        text, page_count = extract_text_from_pdf_bytes(pdf_bytes)
        
        logger.info(f"Extracted {len(text)} characters from {page_count} pages")
        
        return PDFExtractResponse(
            text=text,
            page_count=page_count,
            success=True
        )
        
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return PDFExtractResponse(
            text="",
            page_count=0,
            success=False,
            error=str(e)
        )


@app.post("/extract-pdf-upload", response_model=PDFExtractResponse)
async def extract_pdf_upload(file: UploadFile = File(...)):
    """
    Extract text from an uploaded PDF file.
    
    Alternative endpoint that accepts file upload directly.
    """
    try:
        logger.info(f"Extracting text from uploaded PDF: {file.filename}")
        
        # Read file content
        pdf_bytes = await file.read()
        
        # Extract text
        text, page_count = extract_text_from_pdf_bytes(pdf_bytes)
        
        logger.info(f"Extracted {len(text)} characters from {page_count} pages")
        
        return PDFExtractResponse(
            text=text,
            page_count=page_count,
            success=True
        )
        
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return PDFExtractResponse(
            text="",
            page_count=0,
            success=False,
            error=str(e)
        )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)
